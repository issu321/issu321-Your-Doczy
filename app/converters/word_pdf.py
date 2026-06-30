import os
import subprocess
from pathlib import Path
from io import BytesIO

def convert_word_to_pdf(input_path, output_path):
    """Convert Word document to PDF."""
    # Try multiple methods in order of preference

    # Method 1: pypandoc (requires pandoc and libreoffice)
    try:
        import pypandoc
        pypandoc.convert_file(str(input_path), 'pdf', outputfile=str(output_path))
        return
    except Exception:
        pass

    # Method 2: docx2pdf (Windows only, requires MS Word)
    try:
        from docx2pdf import convert
        convert(str(input_path), str(output_path))
        return
    except Exception:
        pass

    # Method 3: LibreOffice command line
    try:
        cmd = [
            'libreoffice', '--headless', '--convert-to', 'pdf',
            '--outdir', str(Path(output_path).parent),
            str(input_path)
        ]
        subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=60)
        # LibreOffice names output based on input name, so rename if needed
        lo_output = Path(output_path).parent / (Path(input_path).stem + '.pdf')
        if lo_output.exists() and str(lo_output) != str(output_path):
            lo_output.rename(output_path)
        return
    except Exception:
        pass

    # Method 4: Pure Python fallback using reportlab + python-docx
    _convert_word_to_pdf_fallback(input_path, output_path)

def _convert_word_to_pdf_fallback(input_path, output_path):
    """Fallback converter using reportlab."""
    from docx import Document
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY

    doc = Document(input_path)
    pdf = SimpleDocTemplate(str(output_path), pagesize=letter,
                            rightMargin=72, leftMargin=72,
                            topMargin=72, bottomMargin=18)

    styles = getSampleStyleSheet()
    story = []

    # Map docx styles to reportlab styles
    style_map = {
        'Heading 1': styles['Heading1'],
        'Heading 2': styles['Heading2'],
        'Heading 3': styles['Heading3'],
        'Title': styles['Title'],
        'Normal': styles['Normal'],
    }

    for para in doc.paragraphs:
        style = style_map.get(para.style.name, styles['Normal'])
        text = para.text.strip()
        if text:
            story.append(Paragraph(text, style))
            story.append(Spacer(1, 6))

    # Handle tables (basic)
    for table in doc.tables:
        from reportlab.platypus import Table, TableStyle
        from reportlab.lib import colors

        data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            data.append(row_data)

        if data:
            t = Table(data)
            t.setStyle(TableStyle([
                ('GRID', (0, 0), (-1, -1), 1, colors.grey),
                ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
            ]))
            story.append(t)
            story.append(Spacer(1, 12))

    # Handle images (basic)
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            try:
                image = rel.target_part
                image_bytes = image.blob
                img = RLImage(BytesIO(image_bytes), width=4*inch, height=3*inch)
                story.append(img)
                story.append(Spacer(1, 12))
            except Exception:
                pass

    pdf.build(story)
